"""Extracts raw text from uploaded resume files and persists them to disk.

Supports PDF (pdfplumber), DOCX (python-docx), and plain text. Legacy binary
.doc files are rejected with a clear error since there is no lightweight
pure-Python parser for that format among this project's dependencies.
"""

from __future__ import annotations

import io
import uuid
from pathlib import Path

from fastapi import HTTPException, status

from app.config import get_settings

settings = get_settings()

ALLOWED_EXTENSIONS = {"pdf", "doc", "docx", "txt"}
MIN_EXTRACTED_TEXT_LENGTH = 40


def _extension_of(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def validate_resume_file(filename: str, file_bytes: bytes) -> str:
    """Validates extension + size before any storage or LLM work happens.

    Returns the lowercase extension. Raises HTTPException(400) on anything
    invalid so the request fails fast and cleanly.
    """
    if not filename:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No filename provided")

    ext = _extension_of(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type '.{ext}'. Allowed types: {', '.join(sorted(ALLOWED_EXTENSIONS))}.",
        )

    if not file_bytes:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty")

    max_bytes = settings.max_resume_size_mb * 1024 * 1024
    if len(file_bytes) > max_bytes:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File exceeds the maximum allowed size of {settings.max_resume_size_mb}MB",
        )

    return ext


def save_resume_file(candidate_id: uuid.UUID, filename: str, file_bytes: bytes) -> tuple[str, str]:
    """Saves the raw upload to {resume_storage_dir}/{candidate_id}/{uuid}.{ext}.

    Returns (storage_path, file_type).
    """
    ext = _extension_of(filename) or "bin"
    directory = Path(settings.resume_storage_dir) / str(candidate_id)
    directory.mkdir(parents=True, exist_ok=True)
    stored_name = f"{uuid.uuid4()}.{ext}"
    path = directory / stored_name
    path.write_bytes(file_bytes)
    return str(path), ext


def _extract_pdf_text(file_bytes: bytes) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text:
                parts.append(page_text)
    return "\n".join(parts)


def _extract_docx_text(file_bytes: bytes) -> str:
    import docx  # python-docx package, imported as `docx`

    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def _extract_txt_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Extracts raw text from resume bytes.

    Raises ValueError with a clean, user-facing message if the file can't be
    read or yields no usable text (e.g. a scanned/image-only PDF).
    """
    file_type = file_type.lower()
    try:
        if file_type == "pdf":
            text = _extract_pdf_text(file_bytes)
        elif file_type == "docx":
            text = _extract_docx_text(file_bytes)
        elif file_type == "txt":
            text = _extract_txt_text(file_bytes)
        elif file_type == "doc":
            raise ValueError(
                "Legacy .doc files are not supported. Please re-save your resume as "
                ".docx, .pdf, or .txt and upload again."
            )
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    except ValueError:
        raise
    except Exception as exc:  # noqa: BLE001 - surface a clean message for any parser failure
        raise ValueError(f"Could not read the uploaded {file_type.upper()} file: {exc}") from exc

    text = (text or "").strip()
    if len(text) < MIN_EXTRACTED_TEXT_LENGTH:
        raise ValueError(
            "Could not extract readable text from this resume. It may be a scanned "
            "image, corrupted, or empty. Please upload a text-based PDF, DOCX, or TXT file."
        )
    return text
